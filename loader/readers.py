"""各格式 Reader 实现：md/txt/docx/pdf/xlsx。"""
from __future__ import annotations

import os
import sys

from ..base import LoadedDocument
from .base import BaseReader, LoadError


class MarkdownReader(BaseReader):
    suffixes = ("md", "markdown")

    def read(self, path: str, title: str | None = None,
             ocr_images: bool | None = None,
             ocr_backend: str | None = None) -> LoadedDocument:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return LoadedDocument(
            title=title or os.path.basename(path), content=content,
            source_path=path, file_type="md",
        )


class TextReader(BaseReader):
    suffixes = ("txt", "text", "log")

    def read(self, path: str, title: str | None = None,
             ocr_images: bool | None = None,
             ocr_backend: str | None = None) -> LoadedDocument:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return LoadedDocument(
            title=title or os.path.basename(path), content=content,
            source_path=path, file_type="txt",
        )


class DocxReader(BaseReader):
    suffixes = ("docx",)

    def __init__(self, doc_parser_config=None) -> None:
        # heading 修复：通过 PdfDocParserConfig 注入 OCR 参数，集中管理
        self._doc_parser_config = doc_parser_config

    def read(self, path: str, title: str | None = None,
             ocr_images: bool | None = None,
             ocr_backend: str | None = None) -> LoadedDocument:
        # heading 修复：改用 doc_parser.word.v2 解析，输出含 #/## 标题层级的 Markdown
        # 之前用 python-docx 只取段落文本，丢失了表格、图片 OCR、签章等结构信息
        try:
            from ai_sag.doc_parser.word.v2 import parse_word as _parse_word
        except ImportError as e:
            raise LoadError("缺少 doc_parser.word 依赖，请确认安装：pip install python-docx") from e
        kwargs: dict = {}
        if self._doc_parser_config is not None:
            # 请求级 ocr_backend / ocr_images 覆盖 config 默认值（None 时用 config）
            kwargs["ocr_backend"] = ocr_backend if ocr_backend is not None else self._doc_parser_config.ocr_backend
            kwargs["ocr_images"] = ocr_images if ocr_images is not None else self._doc_parser_config.ocr_images
        try:
            result = _parse_word(path, **kwargs)
        except Exception as e:
            # 降级：doc_parser 解析失败时回退到 python-docx 纯文本，保证入库不中断
            try:
                from docx import Document
                doc = Document(path)
                content = "\n".join((p.text or "").strip() for p in doc.paragraphs if (p.text or "").strip())
            except ImportError as ie:
                raise LoadError("缺少 python-docx，请安装：pip install python-docx") from ie
            return LoadedDocument(
                title=title or os.path.basename(path), content=content,
                source_path=path, file_type="docx",
                metadata={"word_parse_fallback": str(e)},
            )
        # heading 修复：剥掉 WordParser 自动生成的元信息头（# 临时文件名 + > 源文件 + > 段落统计 + ---）
        # 否则 chunk.heading 会被临时文件名（如 tmpqvuvlbc8）污染，且统计行会被当正文
        doc_title = title or os.path.splitext(os.path.basename(path))[0]
        content = self._strip_word_meta_header(result.markdown_text or "")
        return LoadedDocument(
            title=doc_title, content=content,
            source_path=path, file_type="docx",
            metadata={
                "paragraph_count": result.paragraph_count,
                "table_count": result.table_count,
                "image_count": result.image_count,
            },
        )

    @staticmethod
    def _strip_word_meta_header(markdown_text: str) -> str:
        """剥掉 WordParser 在 Markdown 开头加的元信息块。

        WordParser 固定输出形如：
            # {file_stem}
            <空行>
            > 源文件: `...`
            > 段落数: N | 表格数: M | 图片数: K
            <空行>
            ---
            <空行>
            {正文}
        这块元信息对检索/抽取无价值，反而会把临时文件名当 H1 污染 chunk.heading。
        本方法定位到第一个 `---` 后的空行，只保留正文部分。
        """
        if not markdown_text:
            return markdown_text
        import re
        # 匹配开头元信息块：# 标题 + 两个 > 引用 + ---，以 --- 结尾
        m = re.match(
            r'^\s*#{1,6}\s+[^\n]*\n+\s*>[^\n]*\n\s*>[^\n]*\n+\s*---\s*\n+',
            markdown_text,
        )
        if m:
            return markdown_text[m.end():].lstrip()
        # 兜底：仅去掉开头的 # 文件名 行（无 --- 分隔符时）
        m = re.match(r'^\s*#{1,6}\s+[^\n]*\n+', markdown_text)
        if m:
            return markdown_text[m.end():].lstrip()
        return markdown_text


class PDFReader(BaseReader):
    suffixes = ("pdf",)

    def __init__(self, doc_parser_config=None) -> None:
        # heading 修复：通过 PdfDocParserConfig 注入 OCR 与 markdown_mode，集中管理
        self._doc_parser_config = doc_parser_config

    def read(self, path: str, title: str | None = None,
             ocr_images: bool | None = None,
             ocr_backend: str | None = None) -> LoadedDocument:
        # heading 修复：改用 doc_parser.pdf.v1 解析，输出含 #/## 标题层级的 Markdown
        # 之前的纯 get_text() 丢失了标题样式信息，导致 chunk.heading 退化为文件名
        try:
            from ai_sag.doc_parser.pdf.v1 import parse_pdf as _parse_pdf
        except ImportError as e:
            raise LoadError("缺少 doc_parser.pdf 依赖，请确认安装：pip install PyMuPDF pymupdf4llm") from e
        # 从 PdfDocParserConfig 读取参数（未配置时用 parse_pdf 默认值）
        kwargs: dict = {}
        if self._doc_parser_config is not None:
            # 请求级 ocr_backend / ocr_images 覆盖 config 默认值（None 时用 config）
            kwargs["ocr_backend"] = ocr_backend if ocr_backend is not None else self._doc_parser_config.ocr_backend
            kwargs["ocr_images"] = ocr_images if ocr_images is not None else self._doc_parser_config.ocr_images
            kwargs["markdown_mode"] = self._doc_parser_config.pdf_markdown_mode
        try:
            result = _parse_pdf(path, **kwargs)
        except Exception as e:
            # 降级：doc_parser 解析失败时回退到纯文本，保证入库不中断
            import fitz
            with fitz.open(path) as d:
                content = "\n\n".join(d[i].get_text() for i in range(d.page_count))
            return LoadedDocument(
                title=title or os.path.basename(path), content=content,
                source_path=path, file_type="pdf",
                metadata={"pdf_parse_fallback": str(e)},
            )
        return LoadedDocument(
            title=title or os.path.basename(path), content=result.markdown_text or "",
            source_path=path, file_type="pdf",
            metadata={"total_pages": result.total_pages, "pdf_type": result.pdf_type},
        )


class ExcelReader(BaseReader):
    """Excel 文档 Reader：使用 doc_parser/excel/v6 解析为 CSV 文本。

    V6 逐行输出 CSV（标准库 csv.writer 处理转义），不做 section 切分，
    避免 V2 markdown 在 SentenceSplitter 下表格行被打散、列名与值分离导致
    人名等实体漏抽（如"汪晨"未被识别为 person）。

    CSV 文本后续由 TableSplitter 按数据行切分，每行以"列名: 值"格式呈现，
    确保表格实体（创建人、需求编号等）列名上下文完整。

    仅支持 .xlsx（OOXML）格式。.xls（旧二进制格式）需经外部转换工具转为 .xlsx 后再入库。
    """

    suffixes = ("xlsx",)

    # 入库防御性校验阈值
    _MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB，超过则拒绝入库
    _XLSX_MAGIC = b"PK\x03\x04"  # xlsx 本质是 zip，开头签名

    def __init__(self, doc_parser_config=None) -> None:
        self._doc_parser_config = doc_parser_config

    def read(self, path: str, title: str | None = None,
             ocr_images: bool | None = None,
             ocr_backend: str | None = None) -> LoadedDocument:
        try:
            import openpyxl  # noqa: F401
        except ImportError as e:
            raise LoadError("缺少 openpyxl，请安装：pip install openpyxl") from e
        self._validate(path)
        from ai_sag.doc_parser.excel.v6.parser import parse_excel as parse_v6
        result = parse_v6(path)
        content = result.to_csv_text()
        return LoadedDocument(
            title=title or os.path.basename(path), content=content,
            source_path=path, file_type="xlsx",
        )

    def _validate(self, path: str) -> None:
        """入库前防御性校验：文件存在性、扩展名、文件大小、文件签名。"""
        # 1. 文件存在性（先于 getsize，避免抛出系统级异常）
        if not os.path.exists(path):
            raise LoadError(f"文件不存在: {path}")

        # 2. 扩展名校验（拦截 .xls 误传为 .xlsx 的情况）
        ext = os.path.splitext(path)[1].lower()
        if ext != ".xlsx":
            raise LoadError(
                f"仅支持 .xlsx 格式；.xls（旧二进制格式）请先转换为 .xlsx：{path}"
            )

        # 3. 文件大小上限
        size = os.path.getsize(path)
        if size > self._MAX_FILE_SIZE:
            raise LoadError(
                f"Excel 文件过大（{size / 1024 / 1024:.1f}MB），"
                f"超过上限 {self._MAX_FILE_SIZE / 1024 / 1024:.0f}MB"
            )
        if size == 0:
            raise LoadError("Excel 文件为空")

        # 4. 文件签名校验（防止后缀欺骗，如 .doc 改名为 .xlsx）
        try:
            with open(path, "rb") as f:
                magic = f.read(4)
            if magic != self._XLSX_MAGIC:
                raise LoadError(
                    f"文件签名不匹配，不是合法的 .xlsx（OOXML/ZIP）文件：{path}"
                )
        except OSError as e:
            raise LoadError(f"读取文件签名失败: {e}") from e


class CSVReader(BaseReader):
    """CSV 文档 Reader：读取 CSV 原始文本，file_type 标记为 csv。

    CSV 本质是表格数据，保留原始 CSV 文本格式（逗号分隔，换行为行分隔），
    由 TableSplitter 用 csv.reader 解析并按数据行切分，每行带表头列名。
    不转 markdown 表格的原因：TableSplitter 内部用 csv.reader 期望逗号分隔的 CSV 文本，
    markdown 表格语法（| 分隔）会导致 csv.reader 解析错乱（整行被当作一列）。
    与 Excel reader 产出格式保持一致（都是 CSV 文本 → TableSplitter）。

    编码优先 utf-8-sig（带 BOM）/utf-8，回退 gbk/gb18030（中文 CSV 常用 gbk）。
    """

    suffixes = ("csv",)

    def read(self, path: str, title: str | None = None,
             ocr_images: bool | None = None,
             ocr_backend: str | None = None) -> LoadedDocument:
        content = self._read_csv(path)
        return LoadedDocument(
            title=title or os.path.basename(path), content=content,
            source_path=path, file_type="csv",
        )

    @staticmethod
    def _read_csv(path: str) -> str:
        """读取 CSV 文件原始文本，保留逗号分隔格式。

        保留原始 CSV 格式供 TableSplitter 的 csv.reader 解析。
        不做 markdown 转换、不做列对齐、不做转义——这些由 TableSplitter 负责。
        """
        for enc in ("utf-8-sig", "utf-8", "gbk", "gb18030"):
            try:
                with open(path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        # 兜底：utf-8 忽略无法解码的字节
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()